import io
import os

import pandas as pd
import PyPDF2
import docx


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_text(file_path: str, file_type: str | None = None) -> str:
    """Extract readable text from a document on disk.

    ``file_type`` should be a lower-case extension (e.g. ``\".pdf\"``). If
    omitted, the extension is taken from ``file_path``.

    Args:
        file_path: Path to the file.
        file_type: Extension including the dot; inferred from ``file_path``
            when omitted.

    Returns:
        Extracted text as a single string.

    Raises:
        ValueError: Unsupported file type.
        FileNotFoundError: File does not exist at ``file_path``.
    """
    ext = (file_type or os.path.splitext(file_path)[1]).lower().strip()

    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    handlers = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".txt": _extract_txt,
        ".csv": _extract_csv,
    }

    if ext not in handlers:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {', '.join(handlers)}"
        )

    return handlers[ext](file_path)


def extract_text_from_bytes(file_bytes: bytes, ext: str) -> str:
    """Extract text from in-memory file bytes (same formats as ``extract_text``)."""
    ext = ext.lower().strip()

    handlers = {
        ".pdf": _extract_pdf_bytes,
        ".docx": _extract_docx_bytes,
        ".txt": _extract_txt_bytes,
        ".csv": _extract_csv_bytes,
    }

    if ext not in handlers:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {', '.join(handlers)}"
        )

    return handlers[ext](file_bytes)


# ---------------------------------------------------------------------------
# Path-based extraction (PyPDF2 / python-docx / open / pandas)
# ---------------------------------------------------------------------------

def _extract_pdf(file_path: str) -> str:
    with open(file_path, "rb") as fh:
        reader = PyPDF2.PdfReader(fh)
        parts: list[str] = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
    return _normalise("\n".join(parts))


def _extract_docx(file_path: str) -> str:
    document = docx.Document(file_path)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    return _normalise("\n".join(paragraphs))


def _extract_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="replace") as fh:
        return _normalise(fh.read())


def _extract_csv(file_path: str) -> str:
    df = pd.read_csv(file_path)
    return _normalise(df.to_string(index=False))


# ---------------------------------------------------------------------------
# Bytes-based helpers
# ---------------------------------------------------------------------------

def _extract_pdf_bytes(data: bytes) -> str:
    reader = PyPDF2.PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return _normalise("\n".join(parts))


def _extract_docx_bytes(data: bytes) -> str:
    document = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    return _normalise("\n".join(paragraphs))


def _extract_txt_bytes(data: bytes) -> str:
    return _normalise(data.decode("utf-8", errors="replace"))


def _extract_csv_bytes(data: bytes) -> str:
    df = pd.read_csv(io.BytesIO(data))
    return _normalise(df.to_string(index=False))


# ---------------------------------------------------------------------------
# Utility
# ---------------------------------------------------------------------------

def _normalise(text: str) -> str:
    """Strip trailing spaces per line and collapse excessive blank lines."""
    lines = text.splitlines()
    result: list[str] = []
    blank_run = 0
    for line in lines:
        stripped = line.rstrip()
        if stripped == "":
            blank_run += 1
            if blank_run <= 1:
                result.append("")
        else:
            blank_run = 0
            result.append(stripped)
    return "\n".join(result).strip()
