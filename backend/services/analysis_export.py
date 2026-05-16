"""Build professional PDF and DOCX exports of analysis results."""

from __future__ import annotations

import io
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ENTITY_SECTION_LABELS: dict[str, str] = {
    "monetary_values": "Monetary values",
    "dates": "Dates",
    "organizations": "Organizations",
    "key_metrics": "Key metrics",
}

_ENTITY_ORDER = ("monetary_values", "dates", "organizations", "key_metrics")


def _get_fpdf_class():
    """Return FPDF class from fpdf2 (required). Legacy ``fpdf`` 1.7.x is not supported."""
    try:
        from fpdf import FPDF
    except ImportError as exc:
        raise ImportError(
            "PDF export requires fpdf2. Install with: pip install fpdf2"
        ) from exc

    import fpdf as fpdf_mod

    version = getattr(fpdf_mod, "__version__", "") or ""
    # Legacy PyPI package "fpdf" is 1.7.x and lacks fpdf2 APIs.
    if version.startswith("1."):
        raise ImportError(
            "Wrong PDF package installed (fpdf 1.7.x). "
            "Uninstall it and install fpdf2: pip uninstall fpdf -y && pip install fpdf2"
        )
    return FPDF


def _safe_filename_base(document_name: str) -> str:
    stem = Path(document_name or "document").stem or "document"
    stem = re.sub(r"[^\w\-. ]", "_", stem, flags=re.UNICODE).strip("._ ") or "document"
    return stem[:120]


def _format_analysis_timestamp(analyzed_at: str) -> str:
    s = (analyzed_at or "").strip()
    if not s:
        from datetime import timezone

        return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    try:
        normalized = s.replace("Z", "+00:00") if s.endswith("Z") else s
        dt = datetime.fromisoformat(normalized)
        if dt.tzinfo is not None:
            return dt.strftime("%Y-%m-%d %H:%M %Z").strip() or dt.strftime("%Y-%m-%d %H:%M UTC")
        return dt.strftime("%Y-%m-%d %H:%M")
    except ValueError:
        return s


def _normalize_entities(entities: dict[str, Any] | None) -> dict[str, list[str]]:
    out: dict[str, list[str]] = {k: [] for k in _ENTITY_ORDER}
    if not isinstance(entities, dict):
        return out
    for k in _ENTITY_ORDER:
        raw = entities.get(k)
        if isinstance(raw, list):
            out[k] = [str(x).strip() for x in raw if x is not None and str(x).strip()]
    return out


def _pdf_safe(text: str) -> str:
    """Core fonts in FPDF only support latin-1."""
    return (text or "").encode("latin-1", errors="replace").decode("latin-1")


def _pdf_cell_line(pdf: Any, w: float, h: float, text: str = "", *, align: str = "") -> None:
    from fpdf.enums import XPos, YPos

    pdf.cell(w, h, text, align=align, new_x=XPos.LMARGIN, new_y=YPos.NEXT)


def _pdf_multi_line(pdf: Any, w: float, h: float, text: str) -> None:
    from fpdf.enums import XPos, YPos

    pdf.multi_cell(w, h, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)


def build_analysis_docx(
    *,
    document_name: str,
    analyzed_at: str,
    summary: str,
    key_points: list[str],
    entities: dict[str, Any] | None,
    analysis_source: str | None = None,
) -> tuple[bytes, str]:
    """Return (file bytes, suggested download filename)."""
    ent = _normalize_entities(entities)
    stamp = _format_analysis_timestamp(analyzed_at)
    base = _safe_filename_base(document_name)
    filename = f"{base}_Analysis_Report.docx"

    doc = Document()
    t = doc.add_heading("Document Analysis Report", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    r = p.add_run("Original document: ")
    r.bold = True
    p.add_run(document_name or "—")

    p2 = doc.add_paragraph()
    r2 = p2.add_run("Analysis date: ")
    r2.bold = True
    p2.add_run(stamp)

    if analysis_source:
        p3 = doc.add_paragraph()
        r3 = p3.add_run("Analysis source: ")
        r3.bold = True
        label = "Gemini AI" if analysis_source.lower() == "gemini" else str(analysis_source)
        p3.add_run(label)

    doc.add_paragraph()

    doc.add_heading("Executive Summary", level=1)
    for line in (summary or "").splitlines() or [""]:
        para = doc.add_paragraph(line)
        for run in para.runs:
            run.font.size = Pt(11)

    doc.add_heading("Key Points", level=1)
    kps = [str(k).strip() for k in (key_points or []) if str(k).strip()]
    if kps:
        for kp in kps:
            doc.add_paragraph(kp, style="List Bullet")
    else:
        doc.add_paragraph("—")

    doc.add_heading("Extracted Entities", level=1)
    for key in _ENTITY_ORDER:
        label = ENTITY_SECTION_LABELS.get(key, key)
        doc.add_heading(label, level=2)
        vals = ent.get(key) or []
        if vals:
            for v in vals:
                doc.add_paragraph(v, style="List Bullet")
        else:
            doc.add_paragraph("None identified")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue(), filename


def build_analysis_pdf(
    *,
    document_name: str,
    analyzed_at: str,
    summary: str,
    key_points: list[str],
    entities: dict[str, Any] | None,
    analysis_source: str | None = None,
) -> tuple[bytes, str]:
    """Return (file bytes, suggested download filename). Requires fpdf2."""
    FPDF = _get_fpdf_class()
    ent = _normalize_entities(entities)
    stamp = _format_analysis_timestamp(analyzed_at)
    base = _safe_filename_base(document_name)
    filename = f"{base}_Analysis_Report.pdf"

    class AnalysisPDF(FPDF):
        def footer(self) -> None:
            self.set_y(-12)
            self.set_font("Helvetica", "I", 8)
            from fpdf.enums import XPos, YPos

            self.cell(
                0,
                8,
                f"Page {self.page_no()}",
                align="C",
                new_x=XPos.RIGHT,
                new_y=YPos.TOP,
            )

    pdf = AnalysisPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_left_margin(18)
    pdf.set_right_margin(18)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 18)
    _pdf_cell_line(pdf, 0, 12, "Document Analysis Report", align="C")
    pdf.ln(2)

    pdf.set_font("Helvetica", "B", 10)
    _pdf_cell_line(pdf, 0, 6, "Original document:")
    pdf.set_font("Helvetica", "", 10)
    _pdf_multi_line(pdf, 0, 6, _pdf_safe(document_name or "-"))

    pdf.set_font("Helvetica", "B", 10)
    _pdf_cell_line(pdf, 0, 6, "Analysis date:")
    pdf.set_font("Helvetica", "", 10)
    _pdf_multi_line(pdf, 0, 6, _pdf_safe(stamp))

    if analysis_source:
        pdf.set_font("Helvetica", "B", 10)
        _pdf_cell_line(pdf, 0, 6, "Analysis source:")
        pdf.set_font("Helvetica", "", 10)
        label = "Gemini AI" if analysis_source.lower() == "gemini" else str(analysis_source)
        _pdf_multi_line(pdf, 0, 6, _pdf_safe(label))

    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 14)
    _pdf_cell_line(pdf, 0, 9, "Executive Summary")
    pdf.set_font("Helvetica", "", 11)
    _pdf_multi_line(pdf, 0, 6, _pdf_safe(summary or "-"))
    pdf.ln(2)

    pdf.set_font("Helvetica", "B", 14)
    _pdf_cell_line(pdf, 0, 9, "Key Points")
    pdf.set_font("Helvetica", "", 11)
    kps = [str(k).strip() for k in (key_points or []) if str(k).strip()]
    if kps:
        for i, kp in enumerate(kps, 1):
            _pdf_multi_line(pdf, 0, 6, _pdf_safe(f"{i}. {kp}"))
    else:
        _pdf_cell_line(pdf, 0, 6, "-")
    pdf.ln(2)

    pdf.set_font("Helvetica", "B", 14)
    _pdf_cell_line(pdf, 0, 9, "Extracted Entities")
    for key in _ENTITY_ORDER:
        label = ENTITY_SECTION_LABELS.get(key, key)
        pdf.set_font("Helvetica", "B", 12)
        _pdf_cell_line(pdf, 0, 8, _pdf_safe(label))
        pdf.set_font("Helvetica", "", 11)
        vals = ent.get(key) or []
        if vals:
            for v in vals:
                _pdf_multi_line(pdf, 0, 6, _pdf_safe(f"  - {v}"))
        else:
            _pdf_multi_line(pdf, 0, 6, "  None identified")
        pdf.ln(1)

    out = pdf.output()
    return bytes(out), filename
